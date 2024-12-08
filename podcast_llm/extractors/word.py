"""
Word document content extractor for podcast generation.

This module provides functionality to extract text content from Word documents
using the python-docx library. It handles reading .docx files and extracting
their text content.

Example:
    >>> from podcast_llm.extractors.word import WordSourceDocument
    >>> extractor = WordSourceDocument('path/to/document.docx')
    >>> content = extractor.extract()
    >>> print(content)
    'The content of the Word document...'

The module supports:
- Modern Word documents (.docx format)
- Extracts text from paragraphs, tables, and other document elements
- Preserves basic text formatting with spaces and line breaks
"""

from pathlib import Path
from typing import Optional
from docx import Document
from podcast_llm.extractors.base import BaseSourceDocument


class WordSourceDocument(BaseSourceDocument):
    """Extracts text content from Word documents using python-docx.

    This class handles extracting text content from .docx files by reading
    through all paragraphs and tables in the document. It preserves basic
    text formatting with appropriate spacing.

    Example:
        >>> extractor = WordSourceDocument('path/to/document.docx')
        >>> content = extractor.extract()
        >>> print(content)
        'The content of the Word document...'

    Attributes:
        src (str): The path to the Word document
        src_type (str): Always 'Word document'
        title (str): A descriptive title combining src_type and source filename
        content (Optional[str]): The extracted document text
    """
    def __init__(self, source: str) -> None:
        self.src = source
        self.src_type = 'Word document'
        self.title = f"{self.src_type}: {Path(source).name}"
        self.content: Optional[str] = None

    def extract(self) -> str:
        """
        Extract text content from the Word document.
        
        Reads through all paragraphs and tables in the document,
        combining their text content with appropriate spacing.
        
        Returns:
            str: The extracted text content
            
        Raises:
            FileNotFoundError: If the source file doesn't exist
            ValueError: If the file is not a valid .docx file
        """
        if not Path(self.src).exists():
            raise FileNotFoundError(f"Word document not found: {self.src}")
            
        if not self.src.endswith('.docx'):
            raise ValueError("File must be a .docx document")
            
        doc = Document(self.src)
        
        # Extract text from paragraphs
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
        
        # Combine all text with appropriate spacing
        all_text = paragraphs + table_text
        self.content = ' '.join(all_text)

        return self.content
