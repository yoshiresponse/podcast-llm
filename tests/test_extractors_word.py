import pytest
from pathlib import Path
from docx import Document
from podcast_llm.extractors.word import WordSourceDocument


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample Word document for testing."""
    doc_path = tmp_path / "test.docx"
    doc = Document()
    
    # Add some paragraphs
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"
    
    doc.save(doc_path)
    return doc_path


def test_word_document_initialization():
    """Test WordSourceDocument initialization."""
    extractor = WordSourceDocument("test.docx")
    assert extractor.src == "test.docx"
    assert extractor.src_type == "Word document"
    assert extractor.title == "Word document: test.docx"
    assert extractor.content is None


def test_extract_content(sample_docx):
    """Test extracting content from a Word document."""
    extractor = WordSourceDocument(str(sample_docx))
    content = extractor.extract()
    
    assert "First paragraph" in content
    assert "Second paragraph" in content
    assert "Cell 1" in content
    assert "Cell 2" in content
    assert "Cell 3" in content
    assert "Cell 4" in content
    assert extractor.content == content


def test_file_not_found():
    """Test handling of non-existent files."""
    extractor = WordSourceDocument("nonexistent.docx")
    with pytest.raises(FileNotFoundError):
        extractor.extract()


def test_invalid_file_extension():
    """Test handling of invalid file extensions."""
    extractor = WordSourceDocument("document.txt")
    with pytest.raises(FileNotFoundError, match="Word document not found: document.txt"):
        extractor.extract()


def test_empty_document(tmp_path):
    """Test handling of empty documents."""
    doc_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(doc_path)
    
    extractor = WordSourceDocument(str(doc_path))
    content = extractor.extract()
    assert content == "" 